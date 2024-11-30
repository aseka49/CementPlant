from db import get_connection
from docx import Document
from datetime import datetime
import os


def execute_query(query, params=None, fetchone=False, commit=False):
    conn = get_connection()
    cursor = conn.cursor()

    try:
        cursor.execute(query, params or ())
        if commit:
            conn.commit()
        return cursor.fetchone() if fetchone else cursor.fetchall()
    except Exception as e:
        return 'Ошибка при выполнении', e

    finally:
        cursor.close()
        conn.close()


# Выполнение запроса с валидацией данных
def execute_query_with_validation(query, params=None, fetchone=False, commit=False):
    try:
        result = execute_query(query, params, fetchone, commit)
        return result
    except Exception as e:
        print(f"Ошибка при выполнении запроса: {e}")
        return None


def validate_and_normalize(quantity, transaction_type, reason):

    if not isinstance(quantity, (int, float)) or quantity <= 0:
        raise ValueError("Ошибка: Количество должно быть положительным числом.")

    validate_transaction_type(transaction_type)

    if not isinstance(reason, str) or not reason.strip():
        raise ValueError("Ошибка: Причина должна быть строкой и не может быть пустой.")

    return quantity, transaction_type, reason


TRANSACTION_TYPES = {
    'purchasing': 'приход',
    'consumption': 'расход'
}


def validate_transaction_type(transaction_type):
    if transaction_type not in TRANSACTION_TYPES.values():
        raise ValueError(f"Ошибка: Некорректный тип транзакции. Возможные значения: {', '.join(TRANSACTION_TYPES.values())}")


def save_ingredients_report(ingredients, report_type="Все ингредиенты"):

    doc = Document()
    doc.add_heading(f'Отчёт: {report_type}', level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = "Название ингредиента"
    header_cells[1].text = "Количество"

    for ingredient_name, stock_amount in ingredients:
        row_cells = table.add_row().cells
        row_cells[0].text = ingredient_name
        row_cells[1].text = f"{stock_amount}"

    # Добавляем дату генерации
    doc.add_paragraph(f"Дата: {date.today().strftime('%d.%m.%Y')}")

    file_path = f"ingredients_report_{report_type.replace(' ', '_')}.docx"
    doc.save(file_path)
    return file_path
