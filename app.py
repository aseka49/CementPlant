import os
from io import BytesIO

from docx import Document
from flask import Flask, render_template, request, send_file, flash, url_for, redirect, session

from household_goods import delete_item, deduct_items, new_item, arrival_items, household_items
from ingredients import arrival_stock, deduct_stock, new_ingredient, delete_ingredient, stock_status, ingredients_table
from recipes import calculate_total_ingredients, delete_recipe, new_recipe, actual_recipes, update_recipe, \
    show_recipe_ingredient, table_recipes
from transactions import transaction_history, items_transaction_history
from utils import execute_query

app = Flask(__name__, template_folder=os.path.join(os.getcwd(), 'templates'))
app.secret_key = 'test123'


@app.route('/')
def index():
    return render_template('index.html')


# Страница с ингредиентами
@app.route('/ingredients')
def ingredients_page():
    ingredients = stock_status()
    return render_template('ingredients.html', ingredients=ingredients)


@app.route('/add_stock', methods=['GET', 'POST'])
def add_stock():
    if request.method == 'POST':
        ingredient_name = request.form['ingredient_name_add']
        quantity = request.form['quantity']
        reason = request.form['reason']
        try:
            quantity = float(quantity)
            if quantity <= 0:
                flash("Ошибка: Количество должно быть положительным числом.", 'danger')
            else:
                arrival_stock(ingredient_name, quantity, reason)
                flash(f"Добавлено {quantity} кг {ingredient_name}.", 'success')
        except ValueError as ve:
            flash(f"Ошибка: {ve}", 'danger')
        except Exception as e:
            flash(f"Произошла непредвиденная ошибка: {e}", 'danger')
        return redirect(url_for('ingredients_page'))
    ingredients = stock_status()
    return render_template('ingredients.html', ingredients=ingredients)


@app.route('/remove_stock', methods=['POST'])
def remove_stock():
    if request.method == 'POST':
        ingredient_name = request.form['ingredient_name_remove']
        quantity = request.form['quantity']
        reason = request.form['reason']
        try:
            quantity = float(quantity)
            if quantity <= 0:
                flash("Ошибка: Количество должно быть положительным числом.", 'danger')
            else:
                success = deduct_stock(ingredient_name, quantity, reason)

                if success:
                    flash(f"Списано {quantity} кг {ingredient_name}.", 'success')
                else:
                    flash(f"Ошибка: Недостаточно {ingredient_name} на складе для списания.", 'danger')
        except ValueError as ve:
            flash(f"Ошибка: {ve}", 'danger')
        except Exception as e:
            flash(f"Произошла непредвиденная ошибка: {e}", 'danger')
        return redirect(url_for('ingredients_page'))

    ingredients = stock_status()
    return render_template('ingredients.html', ingredients=ingredients)


@app.route('/new_ingredient', methods=['POST'])
def new_ingredient_route():
    name = request.form['ingredient_name_new']
    unit = request.form['unit']
    stock_amount = request.form['stock_amount']
    try:
        stock_amount = float(stock_amount)
        result = new_ingredient(name, unit, stock_amount)
        flash(result, 'success')
    except Exception as e:
        flash(f"Произошла ошибка: {e}", 'danger')
    return redirect(url_for('ingredients_page'))

    ingredients = stock_status()
    return render_template('ingredients.html', ingredients=ingredients)


@app.route('/delete_ingredient', methods=['POST'])
def delete_ingredient_route():
    name = request.form['ingredient_name']
    try:
        result = delete_ingredient(name)
        flash(result, 'success')
    except Exception as e:
        flash(f"Произошла ошибка: {e}", 'danger')
    return redirect(url_for('ingredients_page'))

    ingredients = stock_status()
    return render_template('ingredients.html', ingredients=ingredients)


@app.route('/ingredients', methods=['GET', 'POST'])
def ingredients_view():
    sort_type = 'alphabet'  # По умолчанию сортировка по алфавиту
    ingredients = []  # Список для ингредиентов

    if request.method == 'POST':
        # Получаем тип сортировки
        sort_type = request.form.get('sort', 'alphabet')

    ingredients = ingredients_table(sort_type)

    return render_template(
        'ingredients.html',
        ingredients=ingredients,
        current_sort=sort_type
    )


# Страница калькуляции

@app.route('/calculate_ingredients', methods=['GET', 'POST'])
def calculate_ingredients_route():
    if request.method == 'POST':
        recipes = request.form.getlist('recipes[]')
        quantities = request.form.getlist('quantities[]')
        calculation_date = request.form.get('calculation_date')  # Получение введенной даты

        if not calculation_date:
            flash("Ошибка: Дата не выбрана.", 'danger')
            return redirect(url_for('calculate_ingredients_route'))

        recipe_servings = [(recipes[i], int(quantities[i])) for i in range(len(recipes))]

        total_ingredients = calculate_total_ingredients(recipe_servings, calculation_date)

        if total_ingredients is None:
            flash("Ошибка: Недостаточно ингредиентов для приготовления всех блюд.", 'danger')
            return redirect(url_for('calculate_ingredients_route'))

        session['total_ingredients'] = total_ingredients
        session['calculation_date'] = calculation_date

        return render_template('results.html', total_ingredients=total_ingredients, date=calculation_date)

    available_recipes = [row[0] for row in execute_query("SELECT recipe_name FROM recipes ORDER BY recipe_name ASC")]
    return render_template('calculate_ingredients.html', recipes=available_recipes)


@app.route('/download_results', methods=['POST'])
def download_results():
    if 'total_ingredients' not in session or 'calculation_date' not in session:
        return "Ошибка: нет данных для скачивания."

    total_ingredients = session['total_ingredients']
    calculation_date = session['calculation_date']

    document = Document()
    document.add_heading("Общее количество ингредиентов", level=1)
    document.add_paragraph(f"Дата расчета: {calculation_date}")
    document.add_paragraph("Для приготовления всех блюд:")

    for ingredient, data in total_ingredients.items():
        document.add_paragraph(f"{ingredient}: {data['quantity']} кг")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='ingredients_report.docx'
    )


# Страница рецептов

@app.route('/manage_recipes', methods=['GET', 'POST'])
def manage_recipes_page():
    sort_type = request.form.get('sort', 'name')
    recipes = table_recipes(sort_type)
    ingredients = stock_status()
    categories = [row[0] for row in execute_query("SELECT category_name FROM categories")]
    return render_template('manage_recipes.html', recipes=recipes, ingredients=ingredients,categories=categories, current_sort=sort_type)


@app.route('/new_recipe', methods=['POST'])
def new_recipe_route():
    recipe_name = request.form.get('recipe_name')
    ingredient_names = request.form.getlist('ingredients[]')
    quantities = request.form.getlist('quantities[]')
    category_name = request.form.get('category_name')

    try:
        if not recipe_name:
            raise ValueError("Ошибка: Название рецепта не может быть пустым.")

        if not category_name:
            raise ValueError("Ошибка: Выберите категорию для рецепта.")

        # Проверка ингредиентов
        ingredients = []
        for name, qty in zip(ingredient_names, quantities):
            if not name or not qty:
                raise ValueError("Ошибка: Все поля ингредиентов и количеств должны быть заполнены.")
            qty = float(qty)
            if qty <= 0:
                raise ValueError("Ошибка: Количество ингредиента должно быть положительным числом.")
            ingredients.append((name, qty))

        # Попытка добавить рецепт
        message = new_recipe(recipe_name, ingredients, category_name)
        flash(message, 'success')

    except ValueError as ve:
        flash(f"Ошибка валидации: {ve}", 'danger')
    except Exception as e:
        flash(f"Произошла ошибка: {e}", 'danger')

    return redirect(url_for('manage_recipes_page'))  # Перенаправляем на страницу manage_recipes_page




@app.route('/delete_recipe', methods=['POST'])
def delete_recipe_route():
    recipe_name = request.form.get('recipe_name')
    try:
        result = delete_recipe(recipe_name.strip())
        flash(result, 'success')
        return redirect(url_for('manage_recipes_page'))
    except Exception as e:
        flash(f"Произошла ошибка: {e}", 'danger')

    return redirect(url_for('manage_recipes_page'))


# Актуальные рецепты
@app.route('/show_recipes', methods=['GET', 'POST'])
def show_recipes():
    sort_type = request.form.get('sort', 'name') if request.method == 'POST' else 'name'
    recipes = table_recipes(sort_type)

    return render_template('manage_recipes.html', recipes=recipes, current_sort=sort_type)



# Редактирование рецептов

@app.route('/update_recipe')
def recipe_page():
    recipes = actual_recipes()
    return render_template('recipe.html', recipes=recipes)


@app.route('/recipe', methods=['GET', 'POST'])
def recipe():
    if request.method == 'POST':
        recipe_name = request.form.get('recipe_name')

        if 'show_recipe' in request.form:
            ingredients = show_recipe_ingredient(recipe_name)
            return render_template('recipe.html', recipes=actual_recipes(), ingredients=ingredients,
                                   selected_recipe=recipe_name)

        new_recipe_name = request.form.get('new_recipe_name')
        updated_ingredients = []
        removed_ingredients = []

        if request.form.getlist('updated_ingredients[]'):
            for i in range(len(request.form.getlist('updated_ingredients[]'))):
                ingredient_name = request.form.getlist('updated_ingredients[]')[i]
                amount_str = request.form.getlist('updated_amount[]')[i]
                try:
                    amount = float(amount_str) if amount_str else None
                    if amount is not None and amount > 0:
                        updated_ingredients.append((ingredient_name, amount))
                except ValueError:
                    return render_template(
                        'recipe.html', recipes=actual_recipes(),
                        message="Ошибка: Неверный формат количества ингредиента.", selected_recipe=recipe_name
                    )

        removed_ingredients = request.form.getlist('removed_ingredient[]')
        removed_ingredients = [ingredient.strip() for ingredient in removed_ingredients if ingredient.strip()]

        try:
            message = update_recipe(
                recipe_name,
                new_recipe_name=new_recipe_name if new_recipe_name else None,
                updated_ingredients=updated_ingredients,
                removed_ingredients=removed_ingredients
            )
            flash(message,'success')
        except Exception as e:
            message = f"Ошибка при обновлении рецепта: {e}"
            flash(message, 'danger')

        return render_template('recipe.html', recipes=actual_recipes(), message=message, selected_recipe=recipe_name)

    recipes = execute_query("SELECT recipe_name FROM recipes ORDER BY recipe_name ASC", fetchone=False)
    return render_template('recipe.html', recipes=recipes)


# Страница с хозкой

@app.route('/household_items')
def household_items_page():
    items = household_items()
    return render_template('household_items.html', items=items)


@app.route('/add_item', methods=['GET', 'POST'])
def add_item():
    if request.method == 'POST':
        item_name = request.form['item_name']
        quantity = request.form['quantity']
        reason = request.form['reason']
        try:
            quantity = float(quantity)
            if quantity <= 0:
                flash("Ошибка: Количество должно быть положительным числом.", 'danger')

            else:
                arrival_items(item_name, quantity, reason)
                flash(f"Добавлено {quantity} кг {item_name}.", 'success')
        except ValueError as ve:
            flash(f"Ошибка: {ve}", 'danger')
        except Exception as e:
            flash(f"Произошла непредвиденная ошибка: {e}", 'danger')
        return redirect(url_for('household_items_page'))

    items = household_items()
    return render_template('household_items.html', items=items)


@app.route('/remove_item', methods=['POST'])
def remove_item():
    if request.method == 'POST':
        item_name = request.form['item_name']
        quantity = request.form['quantity']
        reason = request.form['reason']
        try:
            quantity = float(quantity)
            if quantity <= 0:
                flash("Ошибка: Количество должно быть положительным числом.", 'danger')
            else:
                succes = deduct_items(item_name,quantity,reason)
                if succes:
                    flash(f"Списано {quantity} шт {item_name}.", 'success')
                else:
                    flash(f"Ошибка: Недостаточно {item_name} для списания", 'danger')
        except ValueError as ve:
            flash(f"Ошибка: {ve}", 'danger')
        except Exception as e:
            flash(f"Произошла непредвиденная ошибка: {e}", 'danger')
        return redirect(url_for('household_items_page'))
    items = household_items()
    return render_template('household_items.html', items=items)


@app.route('/new_item', methods=['POST'])
def new_item_route():
    item_name = request.form['name']
    unit = request.form['unit']
    stock_amount = request.form['stock_amount']
    try:
        stock_amount = float(stock_amount)
        result = new_item(item_name, unit, stock_amount)
        flash(result, 'success')
    except ValueError as ve:
        flash(f"Ошибка: {ve}", 'danger')
    except Exception as e:
        flash(f"Произошла непредвиденная ошибка: {e}", 'danger')
    return redirect(url_for('household_items_page'))

    items = household_items()
    return render_template('household_items.html', items=items)


@app.route('/delete_item_route', methods=['POST'])
def delete_item_route():
    item_name = request.form['item_name']
    try:
        result = delete_item(item_name)
        flash(result, 'success')
        return redirect(url_for('household_items_page'))
    except Exception as e:
        flash(f"Произошла ошибка: {e}", 'danger')
    return redirect(url_for('household_items_page'))


@app.route('/household_items', methods=['GET'])
def show_household_items():
    items = household_items()
    return render_template(url_for('household_items.html'), items=items)


# Страница транзакций

@app.route('/transactions')
def transactions_page():
    return render_template('transactions.html')


@app.route('/transaction_history', methods=['GET', 'POST'])
def transaction_history_view():
    results = None
    if request.method == 'POST':
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')
        transaction_type = request.form.get('type')
        results, file_path = transaction_history(start_date, end_date,transaction_type)
        return render_template('transactions.html', results=results)
    return render_template('transactions.html', results=None)


@app.route('/download_transaction_history', methods=['POST'])
def download_transaction_history():
    return send_file("transaction_history.xlsx", as_attachment=True,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# Транзакции хоз.ки

@app.route('/transaction_history_houseitems', methods=['GET', 'POST'])
def household_transaction_history_view():
    data = None
    if request.method == 'POST':
        start_date = request.form.get('start_date')
        end_date = request.form.get('end_date')
        transaction_type = request.form.get('type')
        data, file_path = items_transaction_history(start_date, end_date, transaction_type)
        return render_template('transactions.html', data=data)
    return render_template('transactions.html', data=None)


@app.route('/download_houseitems_transaction_history', methods=['POST'])
def download_household_transaction_history():
    return send_file("items_transaction_history.xlsx", as_attachment=True,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")




if __name__ == '__main__':
    app.run(debug=True)
